"""Sugeneruoja mažus, teisėtus sintetinius testinius failus (PDF/DOCX/PNG).

Šie failai NĖRA tikri organizacijų dokumentai — jie sukurti šios programos
kūrimo metu tik testavimo tikslams (žr. užduoties reikalavimą "saugok mažus
anoniminius HTML/PDF/DOCX fixture failus arba sukurk teisėtus sintetinius
testinius dokumentus").

Paleidimas: python tests/fixtures/generate_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).resolve().parent


def make_sample_pdf() -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    path = FIXTURES_DIR / "sample.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    text = c.beginText(50, 800)
    text.setFont("Helvetica", 11)
    lines = [
        "SINTETINIS TESTINIS DOKUMENTAS (ne tikras)",
        "",
        "Konkurso pavadinimas: Jaunimo emociniu igudziu stiprinimo programa",
        "Organizatorius: Testine savivaldybe",
        "",
        "1. Bendrosios nuostatos",
        "Siame konkurse paraiskas gali teikti viesosios istaigos, asociacijos",
        "ir kiti juridiniai asmenys, kuriu veikla susijusi su jaunimu.",
        "",
        "2. Finansavimas",
        "Bendras programos biudzetas - 10 000 Eur.",
        "Vieno projekto suma - iki 5 tukst. eurų.",
        "",
        "3. Terminai",
        "Paraiskos priimamos iki rugsejo 15 d.",
        "Veiklos turi buti igyvendintos 2026 m. rugsejo 2-15 d.",
        "",
        "4. Kontaktai",
        "Projekto koordinatore Jone Jonaitiene, tel. 8 686 12345,",
        "el. pastas jone.jonaitiene@testine-savivaldybe.lt",
    ]
    for line in lines:
        text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()
    print(f"Sukurta: {path}")


def make_sample_docx() -> None:
    from docx import Document

    path = FIXTURES_DIR / "sample.docx"
    doc = Document()
    doc.add_heading("SINTETINIS TESTINIS DOKUMENTAS (ne tikras)", level=1)
    doc.add_paragraph("Kvietimas teikti paraiskas: Paaugliu psichikos sveikatos stiprinimo mokymai")
    doc.add_paragraph("Organizatorius: Testines rajono savivaldybes administracija")
    doc.add_heading("Tinkamumo salygos", level=2)
    doc.add_paragraph(
        "Paraiskas teikti gali tik viesosios istaigos, asociacijos ir biudzetines "
        "istaigos, veikiancios jaunimo srityje. Privatūs juridiniai asmenys "
        "(įskaitant mažąsias bendrijas) paraiškų teikti negali, tačiau gali "
        "dalyvauti kaip mokymų ar konsultacijų paslaugų teikėjai pagal atskirą sutartį."
    )
    doc.add_heading("Biudzetas", level=2)
    doc.add_paragraph("Vieno projekto finansavimas - iki 8 000 Eur.")
    doc.add_heading("Kontaktai", level=2)
    doc.add_paragraph(
        "Jaunimo reikalu koordinatorius Petras Petraitis, tel. +370 610 98765, "
        "el. pastas petras.petraitis@testine-rajonas.lt"
    )
    doc.save(str(path))
    print(f"Sukurta: {path}")


def make_scanned_pdf_without_text() -> None:
    """Sukuria PDF, kuriame yra TIK paveikslėlis (imituoja nuskenuotą dokumentą
    be teksto sluoksnio) — jokio realaus teksto objekto PDF struktūroje nėra.
    """
    from PIL import Image, ImageDraw
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    img_path = FIXTURES_DIR / "_scanned_page.png"
    img = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(img)
    # Piešiame stačiakampius, imituojančius teksto eilutes (ne realų tekstą),
    # kad įrodytume, jog PDF turinys yra vaizdas, o ne teksto sluoksnis.
    for i in range(15):
        y = 150 + i * 40
        draw.rectangle([100, y, 900, y + 20], fill="black")
    img.save(img_path)

    path = FIXTURES_DIR / "scanned_no_text.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    c.drawImage(str(img_path), 0, 0, width=A4[0], height=A4[1])
    c.showPage()
    c.save()
    img_path.unlink(missing_ok=True)
    print(f"Sukurta: {path}")


if __name__ == "__main__":
    make_sample_pdf()
    make_sample_docx()
    make_scanned_pdf_without_text()
