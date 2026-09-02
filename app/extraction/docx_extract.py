"""DOCX teksto ištraukimas (python-docx). Senų DOC failų konvertavimui
naudojamas LibreOffice headless (jei įdiegtas Docker konteineryje); jei
konvertuoti nepavyksta, grąžinamas aiškus "nepavyko perskaityti" statusas.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path


@dataclass
class DocxExtractionResult:
    text: str
    method: str  # "python-docx" | "libreoffice_convert" | "failed"
    success: bool


def extract_docx_text(path: str | Path) -> DocxExtractionResult:
    from docx import Document as DocxDocument

    path = Path(path)
    try:
        doc = DocxDocument(str(path))
    except Exception:
        return DocxExtractionResult(text="", method="failed", success=False)

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return DocxExtractionResult(text=text, method="python-docx", success=bool(text.strip()))


def convert_doc_to_docx(path: str | Path, timeout_seconds: int = 60) -> Path | None:
    """Konvertuoja seną .doc į .docx per LibreOffice headless. Grąžina None, jei
    LibreOffice neįdiegtas arba konvertavimas nepavyko (nemetame išimties, kad
    kviečiantis kodas galėtų aiškiai pažymėti "nepavyko perskaityti").
    """
    if shutil.which("soffice") is None:
        return None

    path = Path(path)
    out_dir = Path(tempfile.mkdtemp(prefix="doc2docx_"))
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(out_dir),
                str(path),
            ],
            capture_output=True,
            timeout=timeout_seconds,
            check=True,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None

    converted = out_dir / (path.stem + ".docx")
    return converted if converted.exists() else None
